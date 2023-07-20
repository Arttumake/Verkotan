import pandas as pd
import openpyxl as xl
import math
import glob
import os, time
import shutil
import xlwings as xlw
from datetime import datetime
from docx import Document
from docx.shared import Cm, Pt

# Script to automate the boring parts of the workflow in SAR Lab and to reduce errors overall in reporting
# 2020 summer by Arttu Mäkelä

path = os.getcwd()
doc_name = "Report.docx"
frequencies_excel = "frequency table.xlsx"
excel_files = glob.glob(os.path.join(path, "*.xlsx"))
freq_excel_path = path + "\\" + frequencies_excel
excel_files.remove(freq_excel_path)     # remove excel-document containing frequency table from list of files to be processed
temp_excel = f"{path}\\~${frequencies_excel}"   
if temp_excel in excel_files:    # in case user has frequencey-document open, remove it from list of excels to be processed
    excel_files.remove(temp_excel)
df = pd.read_excel(frequencies_excel, sheet_name='Sheet1')

input_list = df.columns    
print(df.columns[0]) 
freq_list = []
for c in range(df.columns.size):
    print(df[input_list[c]])
    freq_list.append(df[input_list[c]].to_list())
count = 0
for freq in freq_list:  
    freq_list[count] = [x for x in freq if not math.isnan(x)]   # eliminates NaN-values from list
    count += 1
device_freqs = dict(zip(input_list.to_list(), freq_list))   # create a dictionary with Column names being keys and list of frequencies being values
print(device_freqs)
user_inputs = {} 
print("Available bands:")
for x in range(len(input_list)):
    if (x/5).is_integer():
        print("\n")
    print(input_list[x], end ="   ") 
print("\n")
for excel_file in excel_files:
    retry = True
    excel_name = f"{excel_file}"
    while retry == True:    # Asks user to input column names to use for Liquid-Excel
        retry = False
        user_inputs[excel_file] = []    # assign empty list as a value for dictionary, later to hold frequencies
        file_name = excel_file.split(path + "\\")
        print(f"Enter columns for {file_name[1]}:")
        user_input = input()
        parsed_input = user_input.split(",")
        count = 0
        for x in parsed_input:
            parsed_input[count] = x.lstrip(' ')
            x = parsed_input[count]
            count += 1
        for item in parsed_input:
            value_entered = False
            while not value_entered == True:
                for key in device_freqs:
                    if item == key:
                        value_entered = True
                    else:
                        continue
                    break
                if value_entered == False:
                    print("One or more column names were written incorrectly")
                    user_inputs.clear()
                    retry = True
                    break
                elif value_entered == True:
                    for key, value in device_freqs.items():
                        if key == item:
                            user_inputs[excel_file].append(value)
                        else:
                            continue
                        break
excel_freq_amount = {}

# construct the Excel file 
for excel_file in excel_files:
    wb = xl.load_workbook(excel_file)
    ws = wb.active
    ws.title = "Liquid"
    liquid = f"'Liquid'"
    wb.create_sheet("Search", 1)
    ws2 = wb["Search"]
    ws2['A1'] = "Freq"
    ws2['B1'] = "e' target"
    ws2['C1'] = "(S/m) target"
    ws2['D1'] = "e'"
    ws2['E1'] = "(S/m)"
    ws2['F1'] = "e' delta %"
    ws2['G1'] = "(S/m) delta %"

    final_freq_list = []
    for item in user_inputs[excel_file]:
        for value in item:
            final_freq_list.append(value)
    count = 2
    for cell in range(len(final_freq_list)):
        frequency = final_freq_list[count - 2]
        ws2[f'A{count}'] = frequency
        ws2[f'B{count}'] = f'=FORECAST($A{count},OFFSET(INDIRECT("{liquid}!J:J"),MATCH($A{count},INDIRECT("{liquid}!A:A"),1)-1,0,2), OFFSET(INDIRECT("{liquid}!A:A"),MATCH($A{count},INDIRECT("{liquid}!A:A"),1)-1,0,2))'
        ws2[f'C{count}'] = f'=FORECAST($A{count},OFFSET(INDIRECT("{liquid}!L:L"),MATCH($A{count},INDIRECT("{liquid}!A:A"),1)-1,0,2), OFFSET(INDIRECT("{liquid}!A:A"),MATCH($A{count},INDIRECT("{liquid}!A:A"),1)-1,0,2))'
        ws2[f'D{count}'] = f'=FORECAST($A{count},OFFSET(INDIRECT("{liquid}!B:B"),MATCH($A{count},INDIRECT("{liquid}!A:A"),1)-1,0,2), OFFSET(INDIRECT("{liquid}!A:A"),MATCH($A{count},INDIRECT("{liquid}!A:A"),1)-1,0,2))'
        ws2[f'E{count}'] = f'=FORECAST($A{count},OFFSET(INDIRECT("{liquid}!D:D"),MATCH($A{count},INDIRECT("{liquid}!A:A"),1)-1,0,2), OFFSET(INDIRECT("{liquid}!A:A"),MATCH($A{count},INDIRECT("{liquid}!A:A"),1)-1,0,2))'
        ws2[f'F{count}'] = f'=FORECAST($A{count},OFFSET(INDIRECT("{liquid}!S:S"),MATCH($A{count},INDIRECT("{liquid}!A:A"),1)-1,0,2), OFFSET(INDIRECT("{liquid}!A:A"),MATCH($A{count},INDIRECT("{liquid}!A:A"),1)-1,0,2))'
        ws2[f'G{count}'] = f'=FORECAST($A{count},OFFSET(INDIRECT("{liquid}!U:U"),MATCH($A{count},INDIRECT("{liquid}!A:A"),1)-1,0,2), OFFSET(INDIRECT("{liquid}!A:A"),MATCH($A{count},INDIRECT("{liquid}!A:A"),1)-1,0,2))'
        count += 1
    wb.save(excel_file)
    wb.close()
    freqs = count - 2
    excel_freq_amount[excel_file] = freqs

freq_count = 0   
for key, value in excel_freq_amount.items():
    freq_count += value

def save_formulas(excel_file):
    app = xlw.App(visible=False)
    book = app.books.open(excel_file)
    try:
        book.save()
        book.close()
        app.quit()
    except:
        book.close()
        app.quit() 
def update_table():
    document = Document(doc_name)
    table = document.tables[0]
    date_now = datetime.today().strftime(f'%d.%m.%Y')
    for excel_file in excel_files:
        save_formulas(os.path.basename(excel_file))
        wb = xl.load_workbook(excel_file, data_only=True)
        ws = wb['Search']
        row_amount = excel_freq_amount[excel_file]
        file_name = excel_file.split(path + "\\")
        try:
            row_num = 0
            for row in table.rows:
                row_num += 1
            for x in range(row_amount):
                table.add_row()
                first_header = table.rows[0].cells
                header_cells = table.rows[1].cells
                this_row = table.rows[row_num].cells
                this_row[0].text = date_now
                this_row[1].text = "WB Head"
                this_row[2].text = "22"
                str3 = ws.cell(row=x+2, column=1).value
                str4 = round(ws.cell(row=x+2, column=2).value,2)
                str5 = round(ws.cell(row=x+2, column=3).value,2)
                str6 = round(ws.cell(row=x+2, column=4).value,2)
                str7 = round(ws.cell(row=x+2, column=5).value,2)
                str8 = round(ws.cell(row=x+2, column=6).value,2)
                str9 = round(ws.cell(row=x+2, column=7).value,2)
                this_row[3].text = str(str3)
                this_row[4].text = str(str4)
                this_row[5].text = str(str5)
                this_row[6].text = str(str6)
                this_row[7].text = str(str7)
                this_row[8].text = str(str8)
                this_row[9].text = str(str9)
                row_num +=1
        except TypeError:
            print(f"ERROR: Frequencies in {file_name[1]} out of scope")
            time.sleep(5)
            exit()        
        print(f"{file_name[1]} done")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(7)
    headers = header_cells + first_header
    for cell in headers:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font 
                font.size = Pt(8)
                font.bold = True
    document.save(doc_name)
    wb.close()
def create_table():
    document = Document()
    table = document.add_table(rows=freq_count+2, cols=10)
    table.allow_autofit = False
    table.style = document.styles['Table Grid']
    first_header = table.rows[0].cells
    header_cells = table.rows[1].cells
    target_cell = first_header[4].merge(first_header[5])
    measured_cell = first_header[6].merge(first_header[7])
    deviation_cell = first_header[8].merge(first_header[9])
    target_cell.text = "Target"
    measured_cell.text = "Measured"
    deviation_cell.text = "Deviation"
    header_cells[0].text = 'Date'
    header_cells[1].text = 'Tissue Type'
    header_cells[2].text = 'Tissue Temp [\u00B0C]'
    header_cells[3].text = 'Frequency [MHz]'
    header_cells[4].text = "Dielectric Constant [\u03B5] Target"
    header_cells[5].text = "Conductivity \u03C3 [S/m] Target"
    header_cells[6].text = "Dielectric Constant [\u03B5]"
    header_cells[7].text = "Conductivity \u03C3 [S/m]"
    header_cells[8].text = "\u03B5 (%)"
    header_cells[9].text = "\u03C3 (%)"
    date_now = datetime.today().strftime(f'%d.%m.%Y')
    row_num = 2
    for excel_file in excel_files:
        save_formulas(os.path.basename(excel_file))
        wb = xl.load_workbook(excel_file, data_only=True)
        ws = wb['Search']
        file_name = excel_file.split(path + "\\")
        try:
            row_amount = excel_freq_amount[excel_file]
            for x in range(row_amount):
                this_row = table.rows[row_num].cells
                this_row[0].text = date_now
                this_row[1].text = "WB Head"
                this_row[2].text = "22"
                str3 = ws.cell(row=x+2, column=1).value
                str4 = round(ws.cell(row=x+2, column=2).value,2)
                str5 = round(ws.cell(row=x+2, column=3).value,2)
                str6 = round(ws.cell(row=x+2, column=4).value,2)
                str7 = round(ws.cell(row=x+2, column=5).value,2)
                str8 = round(ws.cell(row=x+2, column=6).value,2)
                str9 = round(ws.cell(row=x+2, column=7).value,2)
                this_row[3].text = str(str3)
                this_row[4].text = str(str4)
                this_row[5].text = str(str5)
                this_row[6].text = str(str6)
                this_row[7].text = str(str7)
                this_row[8].text = str(str8)
                this_row[9].text = str(str9)
                row_num +=1
        except TypeError:
            file_name = excel_file.split(path + "\\")
            print(f"ERROR: Frequencies in {file_name[1]} out of scope")
            time.sleep(6)
            exit()
        print(f"{file_name[1]} done")
    headers = header_cells + first_header
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(7)
    for cell in headers:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font 
                font.size = Pt(8)
                font.bold = True
    document.save(doc_name)
    wb.close()

if os.path.exists(doc_name):
    update_table()
else:
    create_table()
#print(input_list)
#print(freq_list)