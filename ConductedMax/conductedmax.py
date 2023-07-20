import pandas as pd
from docx.shared import Pt
from docx import Document

"""Reads csv file with pandas and gets the maximum power(dBm) per band and per channel type, then outputs a
word-document table listing results. Usable with lteresults_cal.txt only"""

doc_name = "conducted maximums.docx"
text_file = "lteresults_cal.txt"
df = pd.read_csv(text_file, sep="\t", engine='python')

column_names = ["Band", "BW", "RBs", "RB Start", "Modulation", "ChanType", "Channel", "Frequency", "Power"]

for index, name in enumerate(column_names): # setting column names in dataframe to correspond to 'column_names' list
    df.columns.values[index] = name

band_list = df["Band"].unique()
bands_and_values = {}   # to hold {band,a list of lists(containing BW, RBs, RB Start etc.)}

for i in range(len(df.index)): # go through each row, create a dictionary with a unique band number and dict value being a list of values
    band_num = df.iloc[i]['Band']
    if band_num in band_list:
        values = [df.iloc[i]['BW'], df.iloc[i]['RBs'], df.iloc[i]['RB Start'], df.iloc[i]['Modulation'],
                  df.iloc[i]['ChanType'], df.iloc[i]['Channel'], df.iloc[i]['Power']]
        if band_num in bands_and_values:
            bands_and_values[int(band_num)].append(values)
        else:
            bands_and_values[int(band_num)] = [values]


max_powers = {}  # dictionary to hold maximum power for each band/channeltype
for band in bands_and_values.keys():    # loop through each band, put low,mid and high channel powers in their own lists
    low_power_list = []
    mid_power_list = []
    high_power_list = []
    for value in bands_and_values[band]:
        if value[4] in ["LOW CH", "Low CH"]:
            low_power_list.append(value)
        elif value[4] in ["MID CH", "Mid CH"]:
            mid_power_list.append(value)
        elif value[4] in ["HIGH CH", "High CH"]:
            high_power_list.append(value)
        else:
            print(f"ERROR: unexpected channel type {value[4]}")
            exit()
    # finds maximum value of each list
    max_low = max(map(lambda x: x[6], low_power_list))
    max_mid = max(map(lambda x: x[6], mid_power_list))
    max_high = max(map(lambda x: x[6], high_power_list))

    # go through low, mid, high channels for each band and make a dict with {band, (max power + other values)}
    for sublist in low_power_list:
        for power in sublist:
            if power == max_low:
                if band in max_powers:
                    max_powers[band].append(sublist)
                else:
                    max_powers[band] = [sublist]
    for sublist in mid_power_list:
        for power in sublist:
            if power == max_mid:
                if band in max_powers:
                    max_powers[band].append(sublist)
                else:
                    max_powers[band] = [sublist]
    for sublist in high_power_list:
        for power in sublist:
            if power == max_high:
                if band in max_powers:
                    max_powers[band].append(sublist)
                else:
                    max_powers[band] = [sublist]

# get total value count in dictionary, to be used in word table for rows
rows = 0
for key, value in max_powers.items():
    for subvalue in value:
        rows += 1

        
"""Creating a word-document with a table containing all the maximum conducted power values. Lists
maximums per band and per channel type
"""
document = Document()
table = document.add_table(rows=rows+1, cols=8)
table.allow_autofit = False
table.style = document.styles['Table Grid']
header_cells = table.rows[0].cells
word_table_headers = ['LTE Band', 'Channel Type', 'Channel #', 'Modulation',
                        'BW', 'RB Size', 'RB Start', 'Power(dBm)']

for num, header in enumerate(word_table_headers):   # fill header cells in table with values in 'word_table_headers' list
    header_cells[num].text = header 

row_num = 0
band_index = 1
# looping through dictionary and filling word table rows with values
for key, value in max_powers.items():
    print(f"Band {key}\n----------------------\n")
    for subvalue in value:  # value = channel type, channel number, modulation etc.
        this_row = table.rows[row_num + 1].cells
        this_row[0].text = str(key)  # band
        this_row[1].text = str(subvalue[4])  # channeltype
        this_row[2].text = str(subvalue[5])  # channelnumber
        this_row[3].text = str(subvalue[3])  # Modulation
        this_row[4].text = str(subvalue[0])  # BW
        this_row[5].text = str(subvalue[1])  # RB Size
        this_row[6].text = str(subvalue[2])  # RB Start
        this_row[7].text = str((round(float(subvalue[6]), 2)))  # Power
        for cell in this_row:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10)
                    if band_index % 2 == 0:  # bold every 2nd band in table
                        font.bold = True
        print(f"{subvalue[4]}: {this_row[7].text} dBm")
        row_num += 1
    print("\n")
    band_index += 1
document.save(doc_name)

