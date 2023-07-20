import pandas as pd
import openpyxl as xl
import xlwings as xlw
import glob
import os
import time
import shutil
from datetime import datetime
from docx import Document
from docx.shared import Cm
from docx.shared import RGBColor

"""
This script automates portions of the SAR data analysis process. SEMCAD data is exported
to a text file and during execution of this script, an excel file is created with data from
all the text files in script folder. The results (which files passed/failed) are shown to the
user in command-line interface and all handled txt files and excel file are moved to a subfolder along
with a log.txt file detailing what was done with the script and when.
Notable class libraries used:
    Pandas class library to read .txt files and write data to excel sheets.
    Openpyxl class library is used to modify cells within the excel file.
    Xlwings to open Excel-document in background (required for accessing formula values)

IMPORTANT FOR USER: currently script attempts to handle all .txt files in its directory. This means that
script directory should only have relevant .txt files (SEMCAD data exports in .txt or .csv format) in it.
SEMCAD data exports also need to have headers enabled in order for the script to extract a name
for each file.

Created by: Arttu Mäkelä
Supervisor: Ilari Kinnunen
"""
start_time = time.time()
path = os.getcwd()
ignore_list = [] # specify the .txt files you don't want to be read by the script here
all_files = glob.glob(os.path.join(path, "*.txt")) # creates a list of all txt files within directory
excel_file_name = 'results.xlsx'

allowed_step_sizes = [0.0075, 0.004, 0.005] # edit this if new step sizes are introduced
number_of_sheets = len(all_files) 
file_dict = {}
step_errors = {}
pass_rate = 0
doc_name = "Report.docx"
log_file_name = "log.txt"
# parsing current time for later use in folder names/reports
current_time = str(time.asctime(time.localtime(time.time())))
folder_time = current_time.replace(":", "-")
sub_dir_parse = folder_time.replace(" ", "_")
sub_dir_name = sub_dir_parse.replace("__", "_")

def create_excel():
    """Reads a list of .txt files and writes them to one excel file, each text file in its own worksheet"""
    try:
        writer = pd.ExcelWriter(excel_file_name) 
        for f in all_files:
            df = pd.read_csv(f, na_values = ['--'] ,skiprows=4,sep="\t\t", engine='python').dropna()    # reads file, skipping first 5 rows and ignoring lines with "--" values  
            df.to_excel(writer, sheet_name=os.path.basename(f))
        writer.save()
    except:     # catches the error and tells the user to include .txt files in directory
        print("Error reading txt files")
        time.sleep(5)
        exit()

def fill_excel_phase_one():
    """ Creating various excel cells/formulas. Excel creation is separated into 3 phases, because some values that appear 
    in excel formulas cannot be accessed unless that excel file is opened (this excel opening in background is done by 
    "open_excel_in_background" function). For example, in phase_two we need variables that are in excel formulas
    created in phase_one."""

    print("Creating Excel sheets")
    # populates excel file and its sheets with various cells
    for f in all_files:
        ws = wb[f'{os.path.basename(f)}'] # assigns the correct worksheet to the variable

        # parsing the text file for certain Grid values for later use in Excel tables
        try:
            get_grids = open(f'{os.path.basename(f)}')  
            lines = get_grids.readlines()
            grid_line = lines[3]
            grids = grid_line.split("Grid: ")
            grid_values = grids[1].split("x")
            get_grids.close()
            grid_x_value = float(grid_values[0])
            grid_y_value = float(grid_values[2])
        except IndexError as error:     # catching an error and telling the user to export data in proper format from SEMCAD
            get_grids.close()
            print(error)
            print(f"Error handling file: {os.path.basename(f)} -> Make all sure SAR field exports are exported with headers and try again")
            time.sleep(6)
            exit()
        # converting some values to floats (initially strings)
        for row in ws.iter_rows(min_row=3, max_col=5):
            for cell in row:                  
                    cell.value = float(cell.value)
        number_of_cells = 0
        for row in ws.iter_rows(min_row=3, max_col=1):
            number_of_cells += 1

        ws['J2'] = "MAX Value of SAR [W/kg]"
        ws['J3'] = f"=MAX(B3:B{number_of_cells + 2})"
        ws['J4'] = "Cell number of MAX value SAR"
        ws['J5'] = "=MATCH(MAX(B:B),B:B,0)"
        ws['G1'] = "X(m)"
        ws['H1'] = "Y(m)"
        ws['G2'] = "Rounded"
        ws['H2'] = "Rounded"
        ws['N2'] = "X(m)"
        ws['O2'] = "Y(m)"
        ws['P2'] = "Z(m)"

        # creating rounded X and Y columns, used for getting step size 
        cell_number = 3
        for row in ws.iter_rows(min_row=3, min_col=7, max_col=7):
            for cell in row:
                cell.value = f"=ROUND(C{cell_number}-C$3,3)+C$3"
                cell_number += 1
        cell_number = 3
        for row in ws.iter_rows(min_row=3, min_col=8, max_col=8):
            for cell in row:
                cell.value = f"=ROUND(D{cell_number}-D$3,3)+D$3"
                cell_number += 1

        ws['N3'] = f"=INDEX(C3:C{number_of_cells + 2},MATCH(MAX(B3:B{number_of_cells + 2}),B3:B{number_of_cells + 2},0))"
        ws['O3'] = f"=INDEX(D3:D{number_of_cells + 2},MATCH(MAX(B3:B{number_of_cells + 2}),B3:B{number_of_cells + 2},0))"
        ws['P3'] = f"=INDEX(E3:E{number_of_cells + 2},MATCH(MAX(B3:B{number_of_cells + 2}),B3:B{number_of_cells + 2},0))"
        ws['J7'] = "Point M2 (Cell value of MAX Value SAR + 1)"
        ws['K3'] = "=(MATCH(MAX(B:B),B:B,0))+1"
        ws['J10'] = "X- & Y-axis zoom scan step size (m)"

        ws['Z2'] = "Is Point next to SAR peak below 3dB?"
        ws['Z5'] = "%-Ratio between m2 and m1: (>=30%)"
        ws['Z8'] = "Re-measurement required?"
        ws['Z11'] = "Minimum distance?"

        ws['AA12'] = "mm"
        ws['Z9'] = '=IF(AND(Z3="No", Z6>=30%), "No", "Yes")'
        
        ws['O9'] = "gx"
        ws['P9'] = "gy"
        ws['O10'] = grid_x_value
        ws['P10'] = grid_y_value
        ws['R1'] = "Lowest measurement points:"
        ws['R2'] = "SAR"
        ws['S2'] = "X"
        ws['T2'] = "Y"
        data_to_grid_ratio = int(number_of_cells/grid_y_value)
        count = 0
        for cells in range(data_to_grid_ratio):
            ws[f'R{3 + count}'] = "=OFFSET($B$3,(ROW()-3)*$P$10,0)"
            ws[f'S{3 + count}'] = "=OFFSET($C$3,(ROW()-3)*$P$10,0)"
            ws[f'T{3 + count}'] = "=OFFSET($D$3,(ROW()-3)*$P$10,0)"
            ws[f'V{3 + count}'] = f"=SQRT((S{3 + count}-$N$3)^2+(T{3 + count}-$O$3)^2)"
            ws[f'W{3 + count}'] = f"=IF(V{3 + count}<($J$11+($J$11/10)),TRUE,FALSE)"
            ws[f'X{3 + count}'] = f'=IF(AND(W{3 + count}=TRUE, (R{3 + count}/$J$3)<=0.501187),"FAIL", "PASS")'
            count += 1
        ws['Z3'] = f'=IF(COUNTIF(X3:X{2 + data_to_grid_ratio},"FAIL"),"Yes","No")'
        ws['AE2'] = "Minimum distance"
        count = 0
        for cells in range(data_to_grid_ratio):
            ws[f'AE{3 + count}'] = f"=ROUND(IF(R{3 + count}/$J$3<0.5, V{count +3}, 0), 4)"
            count += 1
        ws['Z12'] = f"=SMALL(AE3:AE{data_to_grid_ratio + 2},COUNTIF($AE$3:$AE${data_to_grid_ratio + 2},0)+1)*1000"
    wb.save(excel_file_name)

def open_excel_in_background(excel_file_name):
    """Opens the Excel-document in background and then saves/closes it.
    Needed in order to access values in cell formulas later in script"""

    app = xlw.App(visible=False)
    book = app.books.open(excel_file_name)
    try:
        book.save()
        book.close()
        app.quit()
    except:
        book.close()
        app.quit() 

def fill_excel_phase_two():
    """Fills parts of Excel-document, most importantly determining step size value"""
    for f in all_files:
        ws = wb[f'{os.path.basename(f)}']
        ws2 = wb2[f'{os.path.basename(f)}'] 
        point_m1 = ws2['J3'].value
        point_m2 = ws2['K3'].value
        ws['J8'] = f"=B{point_m2}"
        ws['Z6'] = f"=J8/J3 *100"
        
        # calculates step size used in the respective SAR zoom scan
        count = 1
        for rows in ws2.iter_rows(min_row=3, min_col=7, max_col=7):
            for cell in rows:
                if count > 1:
                    current_cell = cell.value
                    prev_cell = ws2.cell(row=(cell.row-1), column =7).value
                    if current_cell != prev_cell:
                        step = abs(prev_cell-current_cell)
                        break
                    else:
                        continue
                    break
                else:
                    count += 1
                    continue
                break
            else:
                continue
            break
        if step > 0.0065 and step < 0.009:
            step_size_rounded = 0.0075
        else:
            step_size_rounded = float(round(step, 4))
        if step_size_rounded not in allowed_step_sizes:
            step_errors[f] = True
        ws['J11'] = step_size_rounded

    wb2.save(excel_file_name)
    wb.save(excel_file_name)
    wb2.close()

def fill_excel_phase_three():
    """Fills finals parts of the Excel-document. Marks the files as Fails, Passes or Errors
    depeding on a string within an excel cell and recolors/organizes the worksheets based on outcome"""

    for f in all_files:
        ws3 = wb3[f"{os.path.basename(f)}"]
        sheet_title = os.path.basename(f)
        ws = wb[f"{os.path.basename(f)}"]
        # parse a particular line from file and use it as file name
        parse_file = open(f'{os.path.basename(f)}')
        lines = parse_file.readlines()
        line_of_filename = lines[2]        # sets 3rd line of .txt file to the variable
            
        split_1 = line_of_filename.split("/Program/")       # two splitting operations to extract a name for .txt file
        split_2 = split_1[1].split("/")     # use the name between "/Program/" and "/" as the 
        sheet_title = split_2[0]
        parse_file.close()

        ws['J13'] = 'filename'
        ws['J14'] = f"{sheet_title}"
        result = ws3['Z9'].value    # sets the value in cell V9 to the variable, either "Yes" or "No"
        os.rename(os.path.basename(f), f"{sheet_title}" + ".txt")
        # checking whether the cell contains the string "Yes" or "No"
        # and classifies the file as "FAIL","PASS" or "ERROR" depending on value
        if result == "No":
            file_dict[f] = "Pass"
        elif result == "Yes":
            file_dict[f] = "Fail"
            ws.sheet_properties.tabColor = '00FF0000'    
        if f in step_errors.keys():
            if step_errors[f] == True:
                file_dict[f] = "Error"
                ws.sheet_properties.tabColor = '00FFFB00'
    wb3.save(excel_file_name)
    wb.save(excel_file_name)
    print("Excel file created successfully")

    wb._sheets.sort(key=lambda ws: ws.title)
    wb3.save(excel_file_name)
    wb.save(excel_file_name)
    wb.close()
    wb3.close()

def print_results():
    """Prints results in CLI showing most important info about the analysed files"""

    fails = 0
    print("Test complete.\n")
    if "Error" in file_dict.values():
        print("The following files have a step size error:")
        print("---------------------------------")
        for key, value in file_dict.items():
            if value == "Error":
                ws = wb[f"{os.path.basename(key)}"]
                filename = ws['J14'].value
                print(filename)  
                fails += 1
        print("---------------------------------")    
    if "Fail" in file_dict.values():
        print("The following files are tagged as 'FAIL':")
        print("---------------------------------")
        for key, value in file_dict.items():
            if value == "Fail":
                ws = wb[f"{os.path.basename(key)}"]
                filename = ws['J14'].value
                print(filename)
                fails += 1
        print("---------------------------------")
        pass_rate = (number_of_sheets - fails)/number_of_sheets * 100
        print(f"{number_of_sheets} files processed with a PASS rate of {'%.1f' % pass_rate} %\n")
    else:
        pass_rate = (number_of_sheets - fails)/number_of_sheets * 100
        print(f"All {number_of_sheets} files are tagged as 'PASS'\n") 

def create_doc():
    """Creates a word-document with a table showing data about each analysed text file"""

    date_now = datetime.today().strftime(f'%d.%m.%Y')
    document = Document()

    document.add_paragraph(date_now)
    table = document.add_table(rows=number_of_sheets+1, cols=5)
    table.style = document.styles['Medium Shading 1 Accent 1']
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Filename'
    header_cells[1].text = 'Horizontal Grid Step [mm]'
    header_cells[2].text = 'Minimum Distance [mm]'
    header_cells[3].text = 'M2/M1 Ratio [%]'
    header_cells[4].text = "Result"
    row = 1
    for key, value in file_dict.items():
        ws = wb[f"{os.path.basename(key)}"]
        ws3 = wb3[f"{os.path.basename(key)}"]
        this_row = table.rows[row].cells
        step_m = ws['J11'].value
        step_mm = float(round(step_m, 4)) * 1000
        m2_m1_ratio = ws3['Z6'].value  
        rounded_ratio = float(round(m2_m1_ratio, 1))
        min_dist = round(ws3['Z12'].value,2)
        this_row[0].text = ws['J14'].value
        this_row[1].text = str(step_mm)
        this_row[2].text = str(min_dist)
        this_row[3].text = str(rounded_ratio)
        if value == "Fail":
            this_row[4].text = "Fail"
        if value == "Pass":
            this_row[4].text = "Pass"
        if value == "Error":
            this_row[4].text = "Error"
        row += 1
        table.autofit = True

    for cell in table.columns[0].cells:
        cell.width = Cm(6)
    document.save(doc_name)

def create_log_file():
    """Creates a log.txt file with statistics about each analysed file.
    Files with step size errors showed on top of list, with fails 2nd and passes 3rd."""

    log_file = open(log_file_name, "w")
    log_file.write(f"{number_of_sheets} files processed with a PASS rate of {'%.1f' % pass_rate} %\r\n")
    log_file.write(f"Date: {current_time} \r\n")
    fail_num = 0
    pass_num = 0
    error_num = 0
    for key, value in file_dict.items():
        if value == "Error":
            ws = wb[f"{os.path.basename(key)}"]
            ws3 = wb3[f"{os.path.basename(key)}"]
            error_num += 1
            sheet_title = f"ERROR {error_num}"
            ws.title = sheet_title
            log_file.write(f"ERROR {error_num}: ")
            filename = ws['J14'].value
            log_file.write(f"{filename} || ")
            step_m = ws['J11'].value
            step_mm = float(round(step_m, 4)) * 1000
            m2_m1_ratio = ws3['Z6'].value  
            rounded_ratio = float(round(m2_m1_ratio, 1)) 
            min_dist = round(ws3['Z12'].value,2)
            log_file.write(f"Step: {step_mm} mm || M2/M1 Ratio: {rounded_ratio}% || Minimum distance: {min_dist} mm\n")
    log_file.write("\n")
    for key, value in file_dict.items():
        if value == "Fail":
            ws = wb[f"{os.path.basename(key)}"]
            ws3 = wb3[f"{os.path.basename(key)}"]
            fail_num += 1
            sheet_title = f"FAIL {fail_num}"
            ws.title = sheet_title
            log_file.write(f"FAIL {fail_num}: ")
            filename = ws['J14'].value
            log_file.write(f"{filename} || ")
            step_m = ws['J11'].value
            step_mm = float(round(step_m, 4)) * 1000
            m2_m1_ratio = ws3['Z6'].value  
            rounded_ratio = float(round(m2_m1_ratio, 1)) 
            min_dist = round(ws3['Z12'].value,2)
            log_file.write(f"Step: {step_mm} mm || M2/M1 Ratio: {rounded_ratio}% || Minimum distance: {min_dist} mm\n")
    log_file.write("\n")
    for key, value in file_dict.items():
        if value == "Pass":
            ws = wb[f"{os.path.basename(key)}"]
            ws3 = wb3[f"{os.path.basename(key)}"]
            pass_num += 1
            sheet_title = f"PASS {pass_num}"
            ws.title = sheet_title
            filename = ws['J14'].value
            log_file.write(f"PASS {pass_num}: ")
            log_file.write(f"{filename} || ")
            step_m = ws['J11'].value
            step_mm = float(round(step_m, 4)) * 1000
            m2_m1_ratio = ws3['Z6'].value  
            rounded_ratio = float(round(m2_m1_ratio, 1)) 
            min_dist = round(ws3['Z12'].value,2)
            log_file.write(f"Step: {step_mm} mm || M2/M1 Ratio: {rounded_ratio}% || Minimum distance: {min_dist} mm\n")
    log_file.close()

def move_files():
    """Moves all files (data files, log.txt, word-doc, excel-file) under dir_name and sub_dir_path."""

    dir_name = 'SEMCAD_data'
    if not os.path.exists(dir_name):
        os.mkdir(dir_name)
        print("Directory" , dir_name , "created")
    sub_dir_path = os.path.join(path,dir_name,sub_dir_name)
    os.mkdir(sub_dir_path)
    print("Moving files to subdirectory...\n")

    log_path = path + "/" + log_file_name
    excel_path = path + "/" + excel_file_name
    doc_path = path + "/" + doc_name
    
    txt_files = glob.glob(os.path.join(path, "*.txt"))
    for txt_file in txt_files:     # move all .txt files in the list to specified subdirectory
        shutil.move(txt_file, sub_dir_path)
    shutil.move(os.path.basename(excel_path), sub_dir_path)
    shutil.move(os.path.basename(doc_path), sub_dir_path)
    #print ('The script took {0} seconds !'.format(time.time() - start_time))

create_excel()

wb = xl.load_workbook(excel_file_name)
fill_excel_phase_one()

open_excel_in_background(excel_file_name)
wb2 = xl.load_workbook(excel_file_name, data_only=True)
fill_excel_phase_two()

open_excel_in_background(excel_file_name)
wb3 = xl.load_workbook(excel_file_name, data_only=True)
fill_excel_phase_three()

print_results()
create_doc()
create_log_file()
move_files()

exit_script = input("Press Enter to exit the script \n")
print("Exiting...")